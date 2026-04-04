"""
Generate sample_legal_contract.docx — a realistic legal document packed with PII
for testing PII Shield anonymization.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.25)

def h1(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    return p

def h2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    return p

def body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p

def ruled():
    doc.add_paragraph("─" * 72)

# ── Title block ───────────────────────────────────────────────────────────────
h1("SETTLEMENT AGREEMENT AND GENERAL RELEASE")
doc.add_paragraph()
body(
    "This Settlement Agreement and General Release (\"Agreement\") is entered into "
    "as of 14 March 2024, by and between:"
)

body(
    "CLAIMANT:  Jonathan R. Hartley, residing at 47 Elmwood Crescent, "
    "Bristol, BS8 2TQ, United Kingdom.  Date of birth: 09 April 1981.  "
    "National Insurance Number: NX 47 23 18 C.  "
    "Email: j.hartley@personalmail.co.uk.  Mobile: +44 7712 334 890."
)

body(
    "RESPONDENT:  Meridian Capital Solutions Ltd, a company incorporated in England "
    "and Wales (Company No. 08834217), whose registered office is at "
    "2nd Floor, 110 Bishopsgate, London, EC2N 4AY.  "
    "Represented by its Chief Legal Officer, Ms. Priya Venkataraman "
    "(priya.venkataraman@meridiancapital.co.uk, +44 20 7946 0312)."
)

ruled()

# ── Recitals ──────────────────────────────────────────────────────────────────
h2("RECITALS")

body(
    "WHEREAS, on or about 3 November 2022, Mr. Hartley commenced employment "
    "with Meridian Capital Solutions Ltd as a Senior Credit Analyst, "
    "Employee Reference MCS-00427, under a contract dated 18 October 2022 "
    "and signed by Head of HR Daniel Kowalczyk "
    "(d.kowalczyk@meridiancapital.co.uk);"
)

body(
    "WHEREAS, on 7 February 2024, Mr. Hartley was summarily dismissed by "
    "his line manager, Thomas J. Osei (t.osei@meridiancapital.co.uk, "
    "direct line +44 20 7946 0388), following an internal disciplinary "
    "investigation (Reference DIS/2024/047);"
)

body(
    "WHEREAS, Mr. Hartley disputes the grounds for dismissal and has filed "
    "an Employment Tribunal claim (Case No. ET/1400782/2024) alleging "
    "unfair dismissal and indirect age discrimination;"
)

body(
    "WHEREAS, both parties wish to resolve all disputes without further "
    "litigation, subject to the terms set out herein;"
)

body("NOW, THEREFORE, in consideration of the mutual promises contained herein, the parties agree as follows:")

ruled()

# ── Payment ───────────────────────────────────────────────────────────────────
h2("1.  SETTLEMENT PAYMENT")

body(
    "1.1  Meridian Capital Solutions Ltd shall pay to Jonathan R. Hartley "
    "a lump-sum settlement of £42,500 (forty-two thousand five hundred pounds), "
    "comprising £18,000 statutory redundancy equivalent and £24,500 "
    "ex-gratia compensation, free of income tax pursuant to s.403 ITEPA 2003."
)

body(
    "1.2  Payment shall be made within 14 days of execution by bank transfer to:"
    "\n        Account Name:   J R Hartley"
    "\n        Sort Code:      20-14-53"
    "\n        Account Number: 63847291"
    "\n        Bank:           Lloyds Bank plc, Bristol City Branch"
    "\n        IBAN:           GB29 LOYD 2014 5363 8472 91"
)

body(
    "1.3  In addition, the Respondent shall settle Mr. Hartley's outstanding "
    "expense claims totalling £1,247.80 (invoice references EXP-2024-0031 "
    "and EXP-2024-0044)."
)

ruled()

# ── Confidentiality ───────────────────────────────────────────────────────────
h2("2.  CONFIDENTIALITY")

body(
    "2.1  Mr. Hartley agrees to keep the fact, terms, and amount of this "
    "Agreement strictly confidential and not to disclose any details to "
    "any person other than his immediate family and his legal adviser, "
    "Ms. Fatima Al-Rashid of Al-Rashid & Partners Solicitors, "
    "12 Queen Square, Bristol, BS1 4NT "
    "(f.alrashid@alrashidpartners.co.uk, +44 117 923 4400)."
)

body(
    "2.2  Meridian Capital Solutions Ltd agrees that no reference to the "
    "disciplinary proceedings shall appear in any employment reference "
    "provided in respect of Mr. Hartley."
)

ruled()

# ── Medical / Occupational Health ─────────────────────────────────────────────
h2("3.  OCCUPATIONAL HEALTH AND MEDICAL MATTERS")

body(
    "3.1  The parties acknowledge that Mr. Hartley was referred to "
    "occupational health on 12 January 2024 following a diagnosis of "
    "generalised anxiety disorder (ICD-10: F41.1) by his GP, "
    "Dr. Sarah Pemberton of Clifton Medical Centre, "
    "9 Clifton Down Road, Bristol, BS8 4AA (Tel: 0117 946 7700)."
)

body(
    "3.2  The Respondent confirms it will contribute £2,000 towards "
    "Mr. Hartley's private therapy costs, payable to "
    "Bristol Wellbeing Clinic (invoice to be addressed to Meridian Capital "
    "Solutions Ltd, Accounts Payable, ref. MED/2024/HRT)."
)

ruled()

# ── Non-disparagement ─────────────────────────────────────────────────────────
h2("4.  NON-DISPARAGEMENT")

body(
    "Both parties agree not to make any public or private disparaging "
    "statements, written or oral, about each other.  For the avoidance of "
    "doubt, this clause covers social-media posts, reviews on platforms "
    "such as Glassdoor and LinkedIn, and communications to industry bodies."
)

ruled()

# ── Reference ─────────────────────────────────────────────────────────────────
h2("5.  EMPLOYMENT REFERENCE")

body(
    "The Respondent shall provide a factual reference confirming that "
    "Jonathan Hartley was employed as Senior Credit Analyst from "
    "3 November 2022 to 7 February 2024, and that he left by mutual "
    "agreement.  All reference requests should be directed to "
    "hr@meridiancapital.co.uk."
)

ruled()

# ── Governing law ─────────────────────────────────────────────────────────────
h2("6.  GOVERNING LAW AND JURISDICTION")

body(
    "This Agreement shall be governed by and construed in accordance with "
    "the laws of England and Wales.  Each party irrevocably submits to the "
    "exclusive jurisdiction of the courts of England and Wales."
)

ruled()

# ── Signatures ────────────────────────────────────────────────────────────────
h2("SIGNATURES")

doc.add_paragraph()
body("Signed by the CLAIMANT:")
body("Name:      Jonathan R. Hartley")
body("Signature: _______________________________")
body("Date:      _______________________________")
body("Witness:   Claire Dunmore, 47 Elmwood Crescent, Bristol, BS8 2TQ")
doc.add_paragraph()
body("Signed for and on behalf of the RESPONDENT:")
body("Name:      Priya Venkataraman")
body("Title:     Chief Legal Officer, Meridian Capital Solutions Ltd")
body("Signature: _______________________________")
body("Date:      _______________________________")
body("Witness:   Oliver Huang, 2nd Floor, 110 Bishopsgate, London, EC2N 4AY")

ruled()

# ── Schedule: Personal data processing notice ─────────────────────────────────
h2("SCHEDULE 1 — DATA PROCESSING NOTICE")

body(
    "For the purposes of UK GDPR, the personal data of Jonathan R. Hartley "
    "(DOB 09/04/1981, NI NX472318C, passport number 513847296 issued "
    "12 June 2019, expiry 12 June 2029) processed in connection with this "
    "Agreement will be retained for a period of 6 years from the date of "
    "execution and then securely destroyed.  Queries should be directed to "
    "the Respondent's Data Protection Officer, "
    "Kevin Nwosu (k.nwosu@meridiancapital.co.uk, +44 20 7946 0299)."
)

# ── Save ──────────────────────────────────────────────────────────────────────
out = os.path.join(os.path.dirname(__file__), "sample_legal_contract.docx")
doc.save(out)
print(f"Saved: {out}")
