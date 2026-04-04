"""
Deanonymization: restore real PII values from placeholder mappings.
"""

from __future__ import annotations

import logging
from pathlib import Path

from .docx_utils import iter_all_wp_elements, replace_across_runs, save_docx

log = logging.getLogger("pii-shield.deanonymizer")
_flog = logging.getLogger("pii-shield.ner")


def deanonymize_text(text: str, mapping: dict) -> str:
    """Replace all placeholders in text with their real values.
    Substitutes longest placeholders first to avoid partial matches.
    """
    for ph in sorted(mapping.keys(), key=len, reverse=True):
        text = text.replace(ph, mapping[ph])
    return text


def deanonymize_docx(docx_path: str | Path, mapping: dict) -> str:
    """Restore placeholders in a .docx, preserving all formatting.
    Returns path to the restored file.
    """
    from docx import Document
    docx_path = Path(docx_path)
    doc = Document(str(docx_path))
    sorted_ph = sorted(mapping.keys(), key=len, reverse=True)
    wns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    _flog.info(f"=== deanonymize_docx | {len(mapping)} placeholders ===")
    for p_elem in iter_all_wp_elements(doc):
        for ph in sorted_ph:
            replace_across_runs(p_elem, ph, mapping[ph], wns)

    out = docx_path.parent / f"{docx_path.stem}_restored.docx"
    save_docx(doc, out)
    _flog.info(f"=== deanonymize_docx done → {out} ===")
    return str(out)


def write_restored_docx(text: str, path: Path):
    """Write plain text to a formatted .docx (used when deanonymizing text output)."""
    from docx import Document
    from docx.shared import Pt, Inches

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    for line in text.split("\n"):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue
        if stripped.isupper() and len(stripped) < 100:
            p = doc.add_paragraph(stripped)
            p.style = doc.styles["Heading 2"]
        elif stripped.startswith("#"):
            level = min(len(stripped) - len(stripped.lstrip("#")), 4)
            p = doc.add_paragraph(stripped.lstrip("# "))
            p.style = doc.styles[f"Heading {level}"]
        else:
            doc.add_paragraph(stripped)

    save_docx(doc, path)
